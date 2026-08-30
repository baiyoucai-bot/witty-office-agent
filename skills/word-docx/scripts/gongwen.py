"""按 GB/T 9704-2012 生成党政机关公文 .docx。

只写输出文件，不改输入。版式判定线见同技能 references/gongwen-format.md。

用沙箱解释器跑：

    <沙箱 Python> gongwen.py --demo --output 样例.docx
    <沙箱 Python> gongwen.py --spec 通知.json --output 通知.docx
    <沙箱 Python> gongwen.py --spec 通知.json --output 通知.docx --font-zhengwen 仿宋_GB2312

spec 是 UTF-8 JSON，字段见 --help-spec。正文层次不用自己标，
写成「一、」「（一）」「1.」「（1）」开头，脚本按序数自动套黑体/楷体/仿宋。

退出码 0 表示写出成功；2 表示 spec 读不进来或字段不合法。
"""

from __future__ import annotations

import argparse
import json
import math
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# GB/T 9704-2012 4.1：A4，天头 37mm±1，订口 28mm±1，版心 156mm×225mm。
# 右边距和下边距标准没写，是由纸张减版心减已知边距算出来的，见 references。
PAGE_W_MM = 210.0
PAGE_H_MM = 297.0
TOP_MM = 37.0
LEFT_MM = 28.0
BANXIN_W_MM = 156.0
BANXIN_H_MM = 225.0
RIGHT_MM = PAGE_W_MM - BANXIN_W_MM - LEFT_MM   # 26.0
BOTTOM_MM = PAGE_H_MM - BANXIN_H_MM - TOP_MM   # 35.0

# 版心 225mm = 637.795pt，每面 22 行并撑满版心 → 每行 28.99pt。
# 不要取常见的 29pt：22×29=638pt 比版心高 0.2pt，Word 会把第 22 行挤到下一页，
# 每面只排 21 行。这里取 28.95pt（579 twips），22 行 = 636.9pt，刚好装下。
LINES_PER_PAGE = 22
CHARS_PER_LINE = 28
LINE_PT = 28.95

# 页码上边缘至版心下边缘 7mm，即距页面下边缘 35-7=28mm。
FOOTER_MM = BOTTOM_MM - 7.0

# 一磅 = 25.4/72 毫米。按字宽算制表位时要用。
PT_MM = 25.4 / 72

# 中文字号对应磅值。公文各要素默认 3 号。
SIZE = {
    "初号": 42.0, "小初": 36.0, "一号": 26.0, "小一": 24.0,
    "二号": 22.0, "小二": 18.0, "三号": 16.0, "小三": 15.0,
    "四号": 14.0, "小四": 12.0, "五号": 10.5, "小五": 9.0,
}

DEFAULT_FONTS = {
    "zhengwen": "仿宋",          # 正文、主送机关、附件说明、署名
    "biaoti": "方正小标宋简体",   # 标题、发文机关标志
    "heiti": "黑体",             # 层次一、密级、紧急程度
    "kaiti": "楷体",             # 层次二、签发人姓名
    "songti": "宋体",            # 页码
}

# GB/T 9704-2012 7.3.3：层次序数决定字体。第三、四层同为仿宋。
LEVELS = (
    ("^[一二三四五六七八九十百]+、", "heiti"),
    ("^（[一二三四五六七八九十百]+）", "kaiti"),
    (r"^\d+[．.]", "zhengwen"),
    (r"^（\d+）", "zhengwen"),
)

SPEC_HELP = """spec JSON 字段（全部可选，缺了就不排该要素）：

版头
  份号        6 位阿拉伯数字，顶格版心左上角第一行
  密级        如「秘密★1年」，顶格第二行
  紧急程度    如「特急」「加急」，顶格密级下一行
  发文机关标志 如「某某集团有限公司文件」，红色小标宋居中
  发文字号    如「某司办〔2026〕15号」
  签发人      上行文才有；有签发人时发文字号左空一字、签发人右空一字

主体
  标题        字符串或字符串数组（数组=手工指定回行）
  主送机关    含末尾全角冒号，如「各分部、各省电力公司：」
  正文        字符串数组，一项一段；以「一、」「（一）」「1.」「（1）」开头的自动套层次字体
  附件说明    字符串或字符串数组，如「1. 巡检清单」
  发文机关署名 如「某某集团有限公司办公厅」
  成文日期    如「2026年8月20日」
  附注        如「（联系人：张某某，010-12345678）」

版记
  抄送        如「各直属单位。」
  印发机关    如「某某集团办公厅」
  印发日期    如「2026年8月20日」

字体        对象，可覆盖 zhengwen / biaoti / heiti / kaiti / songti
"""


def cell_units(text: str) -> float:
    """按字宽算长度：一个汉字算 1，半角字符算 0.5。用于算署名和日期的对齐。"""
    total = 0.0
    for char in text:
        total += 1.0 if ord(char) > 0x2E80 else 0.5
    return total


def char_squeeze_tw(size_pt: float, chars: int = CHARS_PER_LINE) -> int:
    """每行排满 chars 个字所需的字距压缩，单位 1/20 磅（w:spacing 的单位）。

    版心宽 156mm = 442.20pt，28 个 3 号字 = 28×16 = 448pt，天生超 5.80pt。
    不压字距的话每行只能排 27 个字，比标准少一个，而且没有任何报错——
    这是 GB/T 9704 实现里最常见的静默偏差。向下取整保证压够，不会差 0.2pt 又掉回 27。
    """
    need_pt = (BANXIN_W_MM / PT_MM - chars * size_pt) / chars
    return math.floor(need_pt * 20)


def set_font(run, name: str, size_pt: float, *, bold: bool | None = None,
             color: str | None = None, squeeze_tw: int = 0):
    """python-docx 的 font.name 只写 ascii/hAnsi，中文字要自己补 w:eastAsia。

    少了 eastAsia，Word 会拿主题里的默认中文字体去渲染，文档打开是宋体，
    XML 里却明明写着仿宋——这是本技能最容易静默出错的一处，见 references/docx-traps.md。
    """
    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    if squeeze_tw:
        spacing = rpr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            rpr.append(spacing)
        spacing.set(qn("w:val"), str(squeeze_tw))
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _ppr_insert(paragraph, element):
    """pPr 的子元素有固定顺序，直接 append 会排在 jc/ind 后面，成为无效顺序。

    w:pBdr 排在 pStyle/keepNext 之后、shd/spacing/ind/jc 之前，所以往这些之前插。
    """
    ppr = paragraph._p.get_or_add_pPr()
    after = ("w:shd", "w:spacing", "w:ind", "w:jc", "w:textAlignment", "w:rPr", "w:sectPr")
    for tag in after:
        found = ppr.find(qn(tag))
        if found is not None:
            found.addprevious(element)
            return element
    ppr.append(element)
    return element


def fixed_line(paragraph, pt: float = LINE_PT):
    """固定行距。公文要求每面撑满 22 行，只有 EXACTLY 才不被字号带跑。"""
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    return paragraph


def indent_chars(paragraph, *, size_pt: float, left: float = 0, right: float = 0,
                 first: float = 0, hanging: float = 0):
    """按「字」缩进：入参是字数，写进 XML 的是磅（twips = 字数 × 字号 × 20）。

    hanging 给悬挂缩进：首行左移若干字，回行仍对齐 left，用于「附件：」「抄送：」这种
    冒号后要对齐的要素。first 和 hanging 互斥，OOXML 里两者写同一个位置。

    size_pt 必填，不给默认值：四号要素按三号折算会静默偏 14%，见下面第二条实测。

    **不要改成 w:leftChars / w:hangingChars 那套字单位属性。** 名字看着更贴合
    「左空二字」，实测渲染出来却不在标准要求的位置（libreoffice + pdftotext -bbox 量的）：

    | 写法 | 首行 / 回行 |
    |------|------------|
    | `leftChars=500 hangingChars=300` | 5.00 / 8.00 字 —— hanging 加到了回行上，方向反了 |
    | `left=1600 hanging=960`（磅） | **2.00 / 5.00 字** —— 对 |
    | 两者都写 | 5.00 / 8.00 字 —— 字单位优先，补磅救不回来 |
    | `leftChars=100`，run 14pt | 1.15 字 = 1 × 16/14 —— 字宽取 Normal 样式的 16pt，不是本 run 的 |

    所以字单位错两处：悬挂方向反了，字宽取的不是本要素的字号。磅是这里自己按
    size_pt 算的，两处都绕开了。代价是事后在 Word 里改字号缩进不跟着变——公文字号
    由标准定死，不改；真要改字号也得重排每行 28 字，缩进跟不跟着变已经不是重点。
    """
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            jc.addprevious(ind)
        else:
            ppr.append(ind)
    pairs = (
        ("w:left", left),
        ("w:right", right),
        ("w:firstLine", first),
        ("w:hanging", hanging),
    )
    for attr, value in pairs:
        if value:
            ind.set(qn(attr), str(int(round(value * size_pt * 20))))
    return paragraph


def bottom_border(paragraph, *, color: str = "auto", eighths: int = 8):
    """段落下边框当分隔线用。python-docx 没有这个 API，只能手写 XML。

    w:sz 的单位是 1/8 磅：8 = 1 磅，6 = 0.75 磅。
    """
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(eighths))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    _ppr_insert(paragraph, pbdr)
    return paragraph


def page_field(paragraph, font: str, size_pt: float):
    """页码要用 PAGE 域，不能写死数字。python-docx 没有域 API，手写 fldChar。

    公文页码左右各一条一字线，形如 — 1 —。
    """
    lead = set_font(paragraph.add_run("—"), font, size_pt)
    run = paragraph.add_run()
    set_font(run, font, size_pt)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, end):
        run._r.append(node)
    set_font(paragraph.add_run("—"), font, size_pt)
    return lead


def enable_odd_even(document):
    """单双页码分别居右/居左，靠 settings.xml 里的 evenAndOddHeaders 开关。"""
    settings = document.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def setup_layout(document, fonts: dict[str, str]):
    """纸张、页边距、文档网格、页码。"""
    section = document.sections[0]
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    section.top_margin = Mm(TOP_MM)
    section.bottom_margin = Mm(BOTTOM_MM)
    section.left_margin = Mm(LEFT_MM)
    section.right_margin = Mm(RIGHT_MM)
    section.footer_distance = Mm(FOOTER_MM)
    section.header_distance = Mm(TOP_MM)

    # 文档网格：告诉 Word 这一面就是 22 行 × 28 字，由它去压字距。
    # 只设固定行距能保住 22 行，但每行 28 字要靠网格，见 references/gongwen-format.md。
    sect_pr = section._sectPr
    grid = sect_pr.find(qn("w:docGrid"))
    if grid is None:
        grid = OxmlElement("w:docGrid")
        sect_pr.append(grid)
    grid.set(qn("w:type"), "linesAndChars")
    grid.set(qn("w:linePitch"), str(int(round(LINE_PT * 20))))
    grid.set(qn("w:charSpace"), "0")

    enable_odd_even(document)
    for footer, align in (
        (section.footer, WD_ALIGN_PARAGRAPH.RIGHT),
        (section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        para.alignment = align
        fixed_line(para)
        # 单页码居右空一字，双页码居左空一字。
        indent_chars(para, size_pt=SIZE["四号"],
                     right=1 if align == WD_ALIGN_PARAGRAPH.RIGHT else 0,
                     left=1 if align == WD_ALIGN_PARAGRAPH.LEFT else 0)
        page_field(para, fonts["songti"], SIZE["四号"])

    # 正文默认样式一并设成 3 号仿宋，避免遗漏的段落回落到宋体五号。
    style = document.styles["Normal"]
    style.font.name = fonts["zhengwen"]
    style.font.size = Pt(SIZE["三号"])
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), fonts["zhengwen"])
    return section


def right_tab(paragraph, *, chars_from_right: float, size_pt: float):
    """在版心右侧留出若干字宽处放一个靠右制表位。

    w:tab 的 w:pos 是从左页边距量的，不含段落缩进，所以按版心宽减字宽算。
    """
    pos = BANXIN_W_MM - chars_from_right * size_pt * PT_MM
    paragraph.paragraph_format.tab_stops.add_tab_stop(Mm(pos), WD_TAB_ALIGNMENT.RIGHT)
    return paragraph


def para(document, *, align=WD_ALIGN_PARAGRAPH.LEFT):
    item = document.add_paragraph()
    item.alignment = align
    fixed_line(item)
    return item


def blank(document, count: int = 1):
    for _ in range(count):
        para(document)


def level_font(text: str, fonts: dict[str, str]) -> str:
    import re

    for pattern, key in LEVELS:
        if re.match(pattern, text.strip()):
            return fonts[key]
    return fonts["zhengwen"]


def build_banotou(document, spec: dict, fonts: dict[str, str]) -> None:
    """版头：份号、密级、紧急程度、发文机关标志、发文字号、签发人、红色分隔线。"""
    marker = str(spec.get("发文机关标志") or "").strip()
    number = str(spec.get("发文字号") or "").strip()
    signer = str(spec.get("签发人") or "").strip()

    used = 0  # 已占掉的版心行数，用来算发文机关标志那 35mm
    for key, font_key in (("份号", "zhengwen"), ("密级", "heiti"), ("紧急程度", "heiti")):
        value = str(spec.get(key) or "").strip()
        if not value:
            continue
        item = para(document)
        set_font(item.add_run(value), fonts[font_key], SIZE["三号"])
        used += 1

    if marker:
        # 发文机关标志上边缘至版心上边缘 35mm。版心上边缘就是上边距，
        # 所以 35mm 减去份号/密级/紧急程度已占的行高，余量用 space_before 精确顶出来，
        # 不要拿空行凑——空行只能凑到 29pt 的整数倍。
        item = para(document, align=WD_ALIGN_PARAGRAPH.CENTER)
        item.paragraph_format.space_before = Pt(max(0.0, 35.0 / PT_MM - used * LINE_PT))
        set_font(item.add_run(marker), fonts["biaoti"], SIZE["二号"], color="FF0000")
        blank(document)

    if number or signer:
        item = para(document)
        if signer:
            # 有签发人：发文字号左空一字，签发人右空一字，同一行。
            # 靠右制表位把签发人推到版心右侧，不用空格凑。
            indent_chars(item, size_pt=SIZE["三号"], left=1, right=1)
            right_tab(item, chars_from_right=1, size_pt=SIZE["三号"])
            set_font(item.add_run(number), fonts["zhengwen"], SIZE["三号"])
            set_font(item.add_run("\t"), fonts["zhengwen"], SIZE["三号"])
            set_font(item.add_run("签发人："), fonts["zhengwen"], SIZE["三号"])
            set_font(item.add_run(signer), fonts["kaiti"], SIZE["三号"])
        else:
            item.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(item.add_run(number), fonts["zhengwen"], SIZE["三号"])

    if marker or number:
        # 红色分隔线在发文字号之下 4mm，与版心等宽。标准没写线粗，这里取 1 磅。
        # 这一段不放字，行高压到 1pt，否则 4mm 的间距会被 29pt 行高吃掉。
        line = para(document)
        fixed_line(line, pt=1)
        line.paragraph_format.space_before = Pt(4.0 / PT_MM)
        bottom_border(line, color="FF0000", eighths=8)


def build_zhuti(document, spec: dict, fonts: dict[str, str]) -> None:
    """主体：标题、主送机关、正文、附件说明、署名、成文日期、附注。"""
    # 版心宽度撑满 28 字所需的字距压缩。所有 3 号满行要素都要带上，
    # 否则那一段会静默变成每行 27 字。
    squeeze = char_squeeze_tw(SIZE["三号"])
    title = spec.get("标题")
    lines = [title] if isinstance(title, str) else list(title or [])
    lines = [str(line).strip() for line in lines if str(line).strip()]
    if lines:
        # 标题编排于红色分隔线下空二行。
        blank(document, 2)
        for line in lines:
            item = para(document, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_font(item.add_run(line), fonts["biaoti"], SIZE["二号"])

    receiver = str(spec.get("主送机关") or "").strip()
    if receiver:
        # 标题下空一行，主送机关居左顶格。
        blank(document)
        item = para(document)
        set_font(item.add_run(receiver), fonts["zhengwen"], SIZE["三号"], squeeze_tw=squeeze)

    body = spec.get("正文") or []
    if isinstance(body, str):
        body = [body]
    for chunk in body:
        text = str(chunk).strip()
        if not text:
            blank(document)
            continue
        item = para(document, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        indent_chars(item, size_pt=SIZE["三号"], first=2)  # 每自然段左空二字
        set_font(item.add_run(text), level_font(text, fonts), SIZE["三号"], squeeze_tw=squeeze)

    attach = spec.get("附件说明")
    if attach:
        rows = [attach] if isinstance(attach, str) else [str(r).strip() for r in attach if str(r).strip()]
        # 正文下空一行，左空二字编排「附件」二字加全角冒号。
        # 多个附件要加顺序号；回行时与冒号后的首字对齐，所以悬挂缩进 2+3=5 字。
        blank(document)
        for index, row in enumerate(rows):
            item = para(document)
            if len(rows) > 1 and not row[:2].rstrip(".．、").isdigit():
                row = f"{index + 1}.{row}"
            if index == 0:
                # 首行空二字起排「附件：」，回行对齐冒号后首字：左 5 字、首行拉回 3 字。
                indent_chars(item, size_pt=SIZE["三号"], left=5, hanging=3)
                text = f"附件：{row}"
            else:
                indent_chars(item, size_pt=SIZE["三号"], left=5)
                text = row
            set_font(item.add_run(text), fonts["zhengwen"], SIZE["三号"], squeeze_tw=squeeze)

    署名 = str(spec.get("发文机关署名") or "").strip()
    日期 = str(spec.get("成文日期") or "").strip()
    if 署名 or 日期:
        blank(document)
    if 署名:
        item = para(document, align=WD_ALIGN_PARAGRAPH.RIGHT)
        indent_chars(item, size_pt=SIZE["三号"], right=2)
        set_font(item.add_run(署名), fonts["zhengwen"], SIZE["三号"])
    if 日期:
        # 成文日期首字比署名首字右移二字。署名右对齐、右空二字，所以它的首字在
        # 版心宽 − 2 − 署名宽 处，日期左缩进 = 版心宽 − 署名宽 就正好右移二字。
        #
        # 这里刻意**左对齐按左缩进定位**，不用「右对齐 + 右缩进 = 署名宽 − 日期宽」那个
        # 常见写法：后者要知道日期自己有多宽，而日期是「2026年8月20日」这种中西混排，
        # 阿拉伯数字的字宽随目标机器的西文字体变。实测本机渲染成 0.64 字宽左右，不是
        # 想当然的半角 0.5，7 个数字就差出 1 个字——日期只右移了 1 字，不合标准。
        # 左缩进只依赖署名宽（机关名称是纯汉字，1 字就是 1 字），跟日期宽度无关。
        item = para(document, align=WD_ALIGN_PARAGRAPH.LEFT)
        if 署名:
            banxin_chars = BANXIN_W_MM / PT_MM / SIZE["三号"]
            indent_chars(item, size_pt=SIZE["三号"],
                         left=max(0.0, banxin_chars - cell_units(署名)))
        else:
            # 没有署名就没有可对齐的锚，退回右空二字。
            item.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            indent_chars(item, size_pt=SIZE["三号"], right=2)
        set_font(item.add_run(日期), fonts["zhengwen"], SIZE["三号"])

    note = str(spec.get("附注") or "").strip()
    if note:
        # 附注在成文日期下一行，左空二字加圆括号，中间不空行。
        item = para(document)
        indent_chars(item, size_pt=SIZE["三号"], first=2)
        if not note.startswith(("（", "(")):
            note = f"（{note}）"
        set_font(item.add_run(note), fonts["zhengwen"], SIZE["三号"], squeeze_tw=squeeze)


def build_banji(document, spec: dict, fonts: dict[str, str]) -> None:
    """版记：抄送、印发机关和印发日期，三条分隔线首末粗中间细。"""
    copies = str(spec.get("抄送") or "").strip()
    printer = str(spec.get("印发机关") or "").strip()
    printed = str(spec.get("印发日期") or "").strip()
    if not (copies or printer or printed):
        return

    blank(document)
    first = para(document)
    bottom_border(first, eighths=8)  # 首条粗线 0.35mm≈1 磅

    if copies:
        item = para(document)
        # 抄送左右各空一字，回行与冒号后首字对齐，悬挂缩进 1+3=4 字。
        indent_chars(item, size_pt=SIZE["四号"], left=4, right=1, hanging=3)
        set_font(item.add_run(f"抄送：{copies}"), fonts["zhengwen"], SIZE["四号"])
        bottom_border(item, eighths=6)  # 中间细线 0.25mm≈0.75 磅

    if printer or printed:
        item = para(document)
        indent_chars(item, size_pt=SIZE["四号"], left=1, right=1)
        right_tab(item, chars_from_right=1, size_pt=SIZE["四号"])
        set_font(item.add_run(printer), fonts["zhengwen"], SIZE["四号"])
        set_font(item.add_run("\t"), fonts["zhengwen"], SIZE["四号"])
        set_font(item.add_run(printed), fonts["zhengwen"], SIZE["四号"])
        bottom_border(item, eighths=8)  # 末条粗线


DEMO = {
    "份号": "000017",
    "密级": "秘密★1年",
    "紧急程度": "特急",
    "发文机关标志": "某某集团有限公司文件",
    "发文字号": "某司办〔2026〕15号",
    "标题": "关于开展二〇二六年度\n安全生产专项检查的通知",
    "主送机关": "各部门，各分（子）公司：",
    "正文": [
        "为保障年度安全生产平稳有序，经研究，决定在全公司范围内开展安全生产专项检查。现将有关事项通知如下。",
        "一、检查范围",
        "各生产经营场所、重点设备设施，以及近三年发生过安全事件的单位。",
        "（一）生产场所",
        "重点核查消防设施、用电安全和应急通道。",
        "1. 重点设备运行状况逐台记录，不得以「未见异常」概括。",
        "（1）隐患点应拍照留存，标注部位和发现时间。",
        "二、时间安排",
        "自本通知印发之日起至九月三十日止，分自查、抽查、复查三个阶段推进。",
        "三、工作要求",
        "各单位主要负责人对检查结论负责。发现重大隐患应当在两小时内报送公司安全管理部门，不得瞒报、迟报。",
    ],
    "附件说明": ["安全生产专项检查项目清单", "隐患分级判定标准"],
    "发文机关署名": "某某集团有限公司办公厅",
    "成文日期": "2026年8月20日",
    "附注": "（联系人：张某某，010-12345678）",
    "抄送": "公司各部门，各直属单位。",
    "印发机关": "某某集团有限公司办公厅",
    "印发日期": "2026年8月20日",
}


def build(spec: dict, output: Path, fonts: dict[str, str]) -> Path:
    document = Document()
    setup_layout(document, fonts)
    # 默认模板的 body 里只有 sectPr，没有预留空段（实测 len(Document().paragraphs) == 0），
    # 所以份号直接就是版心第一行，不用先清段落。
    build_banotou(document, spec, fonts)
    build_zhuti(document, spec, fonts)
    build_banji(document, spec, fonts)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description="按 GB/T 9704-2012 生成公文 docx")
    parser.add_argument("--spec", help="UTF-8 JSON 公文要素文件")
    parser.add_argument("--demo", action="store_true", help="用内置样例代替 --spec")
    parser.add_argument("--output", required=False, help="输出 .docx 路径")
    parser.add_argument("--help-spec", action="store_true", help="打印 spec 字段说明后退出")
    for key, value in DEFAULT_FONTS.items():
        parser.add_argument(f"--font-{key}", default=None, help=f"覆盖 {key}，默认 {value}")
    args = parser.parse_args()

    if args.help_spec:
        print(SPEC_HELP)
        return 0
    if not args.output:
        print("FAIL 缺 --output", file=sys.stderr)
        return 2
    if not args.spec and not args.demo:
        print("FAIL 要么给 --spec，要么加 --demo", file=sys.stderr)
        return 2

    if args.demo:
        spec = dict(DEMO)
    else:
        path = Path(args.spec)
        try:
            spec = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, ValueError) as exc:
            print(f"FAIL 读不了 spec {path}：{exc}", file=sys.stderr)
            return 2
        if not isinstance(spec, dict):
            print("FAIL spec 顶层必须是 JSON 对象", file=sys.stderr)
            return 2

    fonts = dict(DEFAULT_FONTS)
    fonts.update({k: v for k, v in (spec.get("字体") or {}).items() if k in fonts and v})
    for key in DEFAULT_FONTS:
        override = getattr(args, f"font_{key}", None)
        if override:
            fonts[key] = override

    # 标题里的 \n 当回行，省得调用方一定要传数组。
    title = spec.get("标题")
    if isinstance(title, str) and "\n" in title:
        spec["标题"] = [line for line in title.split("\n") if line.strip()]

    output = build(spec, Path(args.output), fonts)
    print(f"写出：{output}")
    print(f"版式：A4 {PAGE_W_MM:.0f}×{PAGE_H_MM:.0f}mm 版心 {BANXIN_W_MM:.0f}×{BANXIN_H_MM:.0f}mm "
          f"边距 上{TOP_MM:.0f} 下{BOTTOM_MM:.0f} 左{LEFT_MM:.0f} 右{RIGHT_MM:.0f}")
    print(f"行距：固定 {LINE_PT}pt（{LINES_PER_PAGE} 行/面）网格 {CHARS_PER_LINE} 字/行")
    print(f"字体：正文 {fonts['zhengwen']} 标题 {fonts['biaoti']} 层一 {fonts['heiti']} 层二 {fonts['kaiti']}")
    print("下一步：用 check_docx.py 校验，再转 PDF 看图确认。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
